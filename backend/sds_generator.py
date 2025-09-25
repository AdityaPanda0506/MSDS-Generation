# sds_generator.py
# Comprehensive SDS generation module that integrates with sds_data_fetcher.py
# Generates complete Safety Data Sheets and exports to DOCX format

import pandas as pd
import logging
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Import the data fetcher
from sds_data_fetcher import SDSDataFetcher, fetch_compound_data

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class SDSGenerator:
    """
    Comprehensive SDS Generator that uses SDSDataFetcher for data collection
    and generates complete Safety Data Sheets with DOCX export functionality.
    """
    
    def __init__(self):
        self.data_fetcher = SDSDataFetcher()
        self.section_names = {
            1: "Chemical Product and Company Identification",
            2: "Composition and Information on Ingredients", 
            3: "Hazards Identification",
            4: "First Aid Measures",
            5: "Fire-Fighting Measures",
            6: "Accidental Release Measures",
            7: "Handling and Storage",
            8: "Exposure Controls/Personal Protection",
            9: "Physical and Chemical Properties",
            10: "Stability and Reactivity",
            11: "Toxicological Information",
            12: "Ecological Information",
            13: "Disposal Considerations",
            14: "Transport Information",
            15: "Regulatory Information",
            16: "Other Information"
        }
    
    def generate_comprehensive_sds(self, smiles):
        """
        Main method to generate comprehensive SDS from SMILES input.
        Returns complete SDS data structure ready for display or export.
        """
        logger.info(f"[SDS Generator] Starting comprehensive SDS generation for SMILES: {smiles}")
        
        # Fetch all data using the data fetcher
        data = fetch_compound_data(smiles)
        
        if not data or not data.get("basic_data"):
            logger.error("[SDS Generator] Failed to fetch basic compound data")
            return None
        
        # Extract data components
        basic_data = data.get("basic_data", {})
        safety_data = data.get("safety_data", {})
        toxicity_data = data.get("toxicity_data", {})
        physical_properties = data.get("physical_properties", {})
        additional_data = data.get("additional_data", {})
        
        compound_name = basic_data.get("name", "Unknown Compound")
        logger.info(f"[SDS Generator] Generating SDS for: {compound_name}")
        
        # Initialize SDS structure
        sds = {
            f"Section{i}": {
                "title": self.section_names[i],
                "data": {},
                "notes": [],
                "data_sources": []
            } for i in range(1, 17)
        }
        
        # Build each section
        sds["Section1"] = self._build_section_1(basic_data, additional_data)
        sds["Section2"] = self._build_section_2(basic_data, physical_properties, smiles)
        sds["Section3"] = self._build_section_3(basic_data, safety_data, toxicity_data)
        sds["Section4"] = self._build_section_4(safety_data)
        sds["Section5"] = self._build_section_5(safety_data, basic_data)
        sds["Section6"] = self._build_section_6(safety_data)
        sds["Section7"] = self._build_section_7(safety_data)
        sds["Section8"] = self._build_section_8(safety_data)
        sds["Section9"] = self._build_section_9(basic_data, safety_data, physical_properties)
        sds["Section10"] = self._build_section_10(safety_data)
        sds["Section11"] = self._build_section_11(safety_data, toxicity_data)
        sds["Section12"] = self._build_section_12(safety_data, basic_data)
        sds["Section13"] = self._build_section_13(safety_data)
        sds["Section14"] = self._build_section_14(safety_data, basic_data)
        sds["Section15"] = self._build_section_15(safety_data)
        sds["Section16"] = self._build_section_16(data)
        
        logger.info(f"[SDS Generator] SDS generation completed for: {compound_name}")
        return sds
    
    def _build_section_1(self, basic_data, additional_data):
        """Section 1: Chemical Product and Company Identification"""
        echa_data = additional_data.get("echa", {})
        
        section = {
            "title": self.section_names[1],
            "data": {
                "Product Identifier": basic_data.get("name", "Unknown Compound"),
                "Other Names": ", ".join(basic_data.get("synonyms", [])[:3]) if basic_data.get("synonyms") else "Not available",
                "Synonyms": ", ".join(basic_data.get("synonyms", [])[:10]) if basic_data.get("synonyms") else "Not available",
                "CAS Number": basic_data.get("cas", "Not available"),
                "PubChem CID": str(basic_data.get("cid", "Not available")),
                "ECHA Preferred Name": echa_data.get("echa_preferred_name", "Not available"),
                "Molecular Formula": basic_data.get("formula", "Not available"),
                "Molecular Weight": f"{basic_data.get('mw', 'Unknown')} g/mol" if basic_data.get('mw') else "Not available",
                "Product Code": f"CID-{basic_data.get('cid', 'Unknown')}",
                "Date of SDS": datetime.now().strftime("%Y-%m-%d")
            },
            "data_sources": ["PubChem", "ECHA" if echa_data else "Generated"],
            "notes": ["This SDS is generated for research purposes only"]
        }
        return section
    
    def _build_section_2(self, basic_data, physical_properties, smiles):
        """Section 2: Composition and Information on Ingredients"""
        section = {
            "title": self.section_names[2],
            "data": {
                "Chemical Name": basic_data.get("name", "Unknown"),
                "Common Name": basic_data.get("common_name", basic_data.get("name", "Unknown")),
                "CAS Number": basic_data.get("cas", "Not available"),
                "EC Number": "Not assigned",
                "Index Number": "Not assigned",
                "Molecular Formula": basic_data.get("formula", "Not available"),
                "Molecular Weight": f"{basic_data.get('mw', 0):.2f} g/mol" if basic_data.get('mw') else "Not available",
                "SMILES": smiles,
                "InChI Key": "Not available",
                "Concentration/Purity": "≥95% (typical research grade)",
                "Impurities": "May contain trace organic impurities (<5%)",
                "Additives": "None added",
                "Chemical Family": "Organic compound",
                "Hazardous Ingredients": "This entire product",
                "Classification": "Research chemical",
                "Additional Identifiers": {
                    "Hydrogen Bond Donors": physical_properties.get("Hydrogen Bond Donors", "Not available"),
                    "Hydrogen Bond Acceptors": physical_properties.get("Hydrogen Bond Acceptors", "Not available"),
                    "LogP": f"{basic_data.get('logp', 'Unknown')}" if basic_data.get('logp') else "Not available"
                }
            },
            "data_sources": ["PubChem", "RDKit calculations"],
            "notes": ["Composition based on theoretical structure"]
        }
        return section
    
    def _build_section_3(self, basic_data, safety_data, toxicity_data):
        """Section 3: Hazards Identification"""
        hazard_id = safety_data.get("hazard_identification", {})
        
        # Determine hazard level
        toxicity_class = toxicity_data.get("toxicity_class", "Class IV (Low)")
        is_toxic = "Class I" in toxicity_class or "Class II" in toxicity_class
        is_flammable = basic_data.get("logp", 0) > 1.5
        
        # GHS Classification
        ghs_classification = hazard_id.get("GHS Classification", "Not classified")
        if ghs_classification == "Not available" or not ghs_classification:
            if is_toxic and is_flammable:
                ghs_classification = "Acute Tox. 3, Flam. Liq. 3"
            elif is_toxic:
                ghs_classification = "Acute Tox. 3"
            elif is_flammable:
                ghs_classification = "Flam. Liq. 3"
            else:
                ghs_classification = "Not classified"
        
        # Signal Word
        signal_word = hazard_id.get("Signal Word", "")
        if not signal_word or signal_word == "Not available":
            signal_word = "Danger" if is_toxic else "Warning" if is_flammable else "Warning"
        
        # Pictograms
        pictograms = []
        if is_flammable:
            pictograms.append("GHS02 (Flame)")
        if is_toxic:
            pictograms.extend(["GHS06 (Skull and crossbones)", "GHS08 (Health hazard)"])
        if not pictograms:
            pictograms.append("GHS07 (Exclamation mark)")
        
        section = {
            "title": self.section_names[3],
            "data": {
                "GHS Classification": ghs_classification,
                "Signal Word": signal_word,
                "GHS Pictograms": ", ".join(pictograms),
                "Hazard Statements": hazard_id.get("Hazard Statements", "H315 - May cause skin irritation"),
                "Precautionary Statements": hazard_id.get("Precautionary Statements", 
                    "P264 - Wash hands thoroughly after handling. P280 - Wear protective gloves/clothing/eye protection."),
                "Physical Hazards": "Flammable liquid" if is_flammable else "Combustible material",
                "Health Hazards": f"Acute toxicity ({toxicity_class})",
                "Environmental Hazards": "Harmful to aquatic life" if basic_data.get("logp", 0) > 3 else "May cause environmental effects",
                "Routes of Exposure": "Inhalation, Dermal contact, Eye contact, Ingestion",
                "Target Organs": ", ".join(toxicity_data.get("target_organs", ["Not specified"])),
                "Symptoms of Exposure": "Irritation, nausea, dizziness, headache",
                "Medical Conditions Aggravated": "Pre-existing skin, eye, or respiratory conditions",
                "Hazard Class": toxicity_class,
                "Most Important Hazards": f"Toxicity: {toxicity_class}; Flammability: {'Yes' if is_flammable else 'Low'}"
            },
            "data_sources": ["Toxicity predictions", "GHS guidelines", "PubChem data"],
            "notes": ["Classification based on computational predictions"]
        }
        return section
    
    def _build_section_4(self, safety_data):
        """Section 4: First Aid Measures"""
        first_aid = safety_data.get("first_aid", {})
        
        section = {
            "title": self.section_names[4],
            "data": {
                "General": first_aid.get("General First Aid", 
                    "Remove from exposure immediately. Get medical attention if symptoms persist."),
                "Inhalation": first_aid.get("Inhalation", 
                    "Move to fresh air immediately. If breathing is difficult, give oxygen. Seek medical attention."),
                "Skin Contact": first_aid.get("Skin Contact", 
                    "Remove contaminated clothing. Wash with soap and water for at least 15 minutes. Seek medical attention if irritation persists."),
                "Eye Contact": first_aid.get("Eye Contact", 
                    "Flush immediately with water for at least 15 minutes. Remove contact lenses if present. Seek immediate medical attention."),
                "Ingestion": first_aid.get("Ingestion", 
                    "Rinse mouth with water. Do not induce vomiting. Give water to drink if conscious. Seek immediate medical attention."),
                "Most Important Symptoms": first_aid.get("Most Important Symptoms", 
                    "Irritation of skin, eyes, and respiratory tract. Nausea, dizziness."),
                "Immediate Medical Attention": "Required for significant exposures or persistent symptoms",
                "Notes to Physician": first_aid.get("Notes to Physician", 
                    "Treat symptomatically. Show this SDS to medical personnel."),
                "Specific Treatment": "No specific antidote. Provide supportive care.",
                "Protection of First Aiders": "Use appropriate protective equipment to avoid exposure."
            },
            "data_sources": ["PubChem safety data", "Standard first aid protocols"],
            "notes": ["Follow standard chemical exposure first aid procedures"]
        }
        return section
    
    def _build_section_5(self, safety_data, basic_data):
        """Section 5: Fire-Fighting Measures"""
        fire_data = safety_data.get("fire_fighting", {})
        physical_props = safety_data.get("physical_properties", {})
        
        is_flammable = basic_data.get("logp", 0) > 1.5
        
        section = {
            "title": self.section_names[5],
            "data": {
                "Flash Point": physical_props.get("Flash Point", 
                    "< 23°C (predicted)" if is_flammable else "> 93°C (predicted)"),
                "Auto-ignition Temperature": physical_props.get("Auto-ignition Temperature", "Not determined"),
                "Flammable Limits (LEL/UEL)": f"LEL: {physical_props.get('Lower Explosive Limit', 'Not determined')} | UEL: {physical_props.get('Upper Explosive Limit', 'Not determined')}",
                "Suitable Extinguishing Media": fire_data.get("Extinguishing Media", 
                    "Carbon dioxide, dry chemical powder, alcohol-resistant foam, water spray"),
                "Unsuitable Extinguishing Media": fire_data.get("Unsuitable Extinguishing Media", 
                    "Water jet (may spread fire)"),
                "Special Fire Fighting Procedures": "Use water spray to cool containers. Fight fire from upwind position.",
                "Protective Equipment for Firefighters": fire_data.get("Special Protective Equipment", 
                    "Self-contained breathing apparatus (SCBA) and full protective clothing"),
                "Unusual Fire/Explosion Hazards": fire_data.get("Special Hazards", 
                    "Vapors may form explosive mixtures with air. Vapors heavier than air."),
                "Hazardous Combustion Products": fire_data.get("Hazardous Combustion Products", 
                    "Carbon monoxide, carbon dioxide, nitrogen oxides, toxic organic compounds"),
                "Fire Fighting Measures": "Evacuate area. Use appropriate extinguishing media.",
                "Sensitivity to Static Discharge": "May be sensitive" if is_flammable else "Not sensitive"
            },
            "data_sources": ["Fire safety guidelines", "Chemical property predictions"],
            "notes": ["Fire fighting procedures based on chemical class"]
        }
        return section
    
    def _build_section_6(self, safety_data):
        """Section 6: Accidental Release Measures"""
        release_data = safety_data.get("accidental_release", {})
        
        section = {
            "title": self.section_names[6],
            "data": {
                "Personal Precautions": release_data.get("Personal Precautions", 
                    "Evacuate personnel. Wear appropriate PPE. Ensure adequate ventilation. Eliminate ignition sources."),
                "Environmental Precautions": release_data.get("Environmental Precautions", 
                    "Prevent entry into waterways, sewers, or soil. Contain spill to minimize environmental impact."),
                "Methods of Containment": release_data.get("Methods of Containment", 
                    "Stop leak if safe to do so. Contain with non-combustible absorbent material."),
                "Methods of Cleaning Up": release_data.get("Methods of Cleaning Up", 
                    "Absorb with inert material. Collect in appropriate containers for disposal."),
                "Small Spills": "Absorb with paper towels or cloth. Dispose as chemical waste.",
                "Large Spills": "Evacuate area. Use appropriate absorbents. Prevent environmental release.",
                "Equipment Needed": "Absorbent materials, non-sparking tools, appropriate containers",
                "Emergency Procedures": "Follow emergency response plan. Notify authorities if required.",
                "Reference to Other Sections": "See Sections 8 and 13 for exposure controls and disposal"
            },
            "data_sources": ["Spill response guidelines"],
            "notes": ["Follow institutional spill response procedures"]
        }
        return section
    
    def _build_section_7(self, safety_data):
        """Section 7: Handling and Storage"""
        handling_data = safety_data.get("handling_storage", {})
        
        section = {
            "title": self.section_names[7],
            "data": {
                "Precautions for Safe Handling": handling_data.get("Precautions for Safe Handling", 
                    "Use in well-ventilated areas. Avoid contact with skin and eyes. Ground containers when transferring."),
                "Conditions for Safe Storage": handling_data.get("Conditions for Safe Storage", 
                    "Store in cool, dry place. Keep container tightly closed. Store away from incompatible materials."),
                "Storage Temperature": handling_data.get("Storage Temperature", "Room temperature (15-25°C)"),
                "Incompatible Materials": handling_data.get("Incompatible Materials", 
                    "Strong oxidizing agents, strong acids, strong bases"),
                "Container Materials": "Glass, PTFE, stainless steel. Avoid reactive metals.",
                "Storage Requirements": "Secondary containment recommended. Proper labeling required.",
                "Shelf Life": "Use within recommended timeframe. Check for degradation.",
                "Special Precautions": "Secure against unauthorized access. Follow local regulations.",
                "Handling Equipment": "Use appropriate tools and containers. Ground equipment when transferring."
            },
            "data_sources": ["Chemical storage guidelines"],
            "notes": ["Follow institutional chemical storage procedures"]
        }
        return section
    
    def _build_section_8(self, safety_data):
        """Section 8: Exposure Controls/Personal Protection"""
        exposure_data = safety_data.get("exposure_controls", {})
        
        section = {
            "title": self.section_names[8],
            "data": {
                "Occupational Exposure Limits": {
                    "TLV-TWA": exposure_data.get("TLV-TWA", "Not established"),
                    "TLV-STEL": exposure_data.get("TLV-STEL", "Not established"),
                    "PEL": exposure_data.get("PEL", "Not established"),
                    "IDLH": exposure_data.get("IDLH", "Not determined")
                },
                "Engineering Controls": exposure_data.get("Engineering Controls", 
                    "Local exhaust ventilation, fume hoods, adequate general ventilation"),
                "Personal Protective Equipment": {
                    "Eye Protection": exposure_data.get("Eye Protection", "Safety goggles or face shield"),
                    "Skin Protection": exposure_data.get("Skin Protection", 
                        "Chemical-resistant gloves (nitrile, neoprene). Lab coat, long pants."),
                    "Respiratory Protection": exposure_data.get("Respiratory Protection", 
                        "Use in well-ventilated area. Respirator if ventilation inadequate."),
                    "Foot Protection": "Closed-toe shoes. Chemical-resistant boots for large quantities."
                },
                "Thermal Hazards": exposure_data.get("Thermal Hazards", "Not applicable at room temperature"),
                "Hygiene Measures": "Wash hands thoroughly after handling. No eating, drinking, or smoking in work areas.",
                "Environmental Controls": "Prevent release to environment. Use appropriate containment."
            },
            "data_sources": ["Exposure control guidelines"],
            "notes": ["Adjust PPE based on quantity and exposure potential"]
        }
        return section
    
    def _build_section_9(self, basic_data, safety_data, physical_properties):
        """Section 9: Physical and Chemical Properties"""
        physical_props = safety_data.get("physical_properties", {})
        mw = basic_data.get("mw", 300)
        
        # Predict physical state based on molecular weight and structure
        predicted_state = "Liquid" if mw < 300 else "Solid"
        predicted_appearance = "Clear liquid" if mw < 300 else "White to off-white solid"
        
        section = {
            "title": self.section_names[9],
            "data": {
                "Physical State": physical_props.get("Physical State", predicted_state),
                "Appearance": physical_props.get("Appearance", predicted_appearance),
                "Color": physical_props.get("Color", "Colorless to pale yellow"),
                "Odor": physical_props.get("Odor", "Characteristic organic odor"),
                "Odor Threshold": physical_props.get("Odor Threshold", "Not determined"),
                "pH": physical_props.get("pH", "Not applicable"),
                "Melting Point": physical_props.get("Melting Point", "Not determined"),
                "Boiling Point": physical_props.get("Boiling Point", "Not determined"),
                "Flash Point": physical_props.get("Flash Point", "Not determined"),
                "Evaporation Rate": physical_props.get("Evaporation Rate", "Not determined"),
                "Flammability": physical_props.get("Flammability", "Combustible"),
                "Upper/Lower Explosive Limits": f"{physical_props.get('Upper Explosive Limit', 'ND')} / {physical_props.get('Lower Explosive Limit', 'ND')}",
                "Vapor Pressure": physical_props.get("Vapor Pressure", "Not determined"),
                "Vapor Density": physical_props.get("Vapor Density", "Not determined"),
                "Density": physical_props.get("Density", "Not determined"),
                "Relative Density": physical_props.get("Relative Density", "Not determined"),
                "Solubility in Water": basic_data.get("solubility", physical_props.get("Solubility in Water", "Not determined")),
                "Partition Coefficient": f"log P = {basic_data.get('logp', 'Not determined')}",
                "Auto-ignition Temperature": physical_props.get("Auto-ignition Temperature", "Not determined"),
                "Decomposition Temperature": physical_props.get("Decomposition Temperature", "Not determined"),
                "Viscosity": physical_props.get("Kinematic Viscosity", "Not determined"),
                # Add computed properties
                "Molecular Weight": f"{mw:.2f} g/mol" if mw else "Not available",
                "Molecular Formula": basic_data.get("formula", "Not available"),
                "Heavy Atom Count": physical_properties.get("Heavy Atom Count", "Not available"),
                "Hydrogen Bond Donors": physical_properties.get("Hydrogen Bond Donors", "Not available"),
                "Hydrogen Bond Acceptors": physical_properties.get("Hydrogen Bond Acceptors", "Not available"),
                "Rotatable Bonds": physical_properties.get("Rotatable Bonds", "Not available"),
                "TPSA": physical_properties.get("Topological Polar Surface Area (TPSA)", "Not available")
            },
            "data_sources": ["RDKit calculations", "Property predictions"],
            "notes": ["Physical properties estimated from molecular structure"]
        }
        return section
    
    def _build_section_10(self, safety_data):
        """Section 10: Stability and Reactivity"""
        stability_data = safety_data.get("stability_reactivity", {})
        
        section = {
            "title": self.section_names[10],
            "data": {
                "Reactivity": stability_data.get("Reactivity", "May be reactive under certain conditions"),
                "Chemical Stability": stability_data.get("Chemical Stability", "Stable under recommended storage conditions"),
                "Possibility of Hazardous Reactions": stability_data.get("Possibility of Hazardous Reactions", 
                    "None under normal storage and handling conditions"),
                "Conditions to Avoid": stability_data.get("Conditions to Avoid", 
                    "Heat, sparks, open flames, strong oxidizing agents"),
                "Incompatible Materials": stability_data.get("Incompatible Materials", 
                    "Strong acids, strong bases, strong oxidizing agents, reactive metals"),
                "Hazardous Decomposition Products": stability_data.get("Hazardous Decomposition", 
                    "Carbon monoxide, carbon dioxide, nitrogen oxides, toxic organic compounds"),
                "Hazardous Polymerization": stability_data.get("Hazardous Polymerization", "Will not occur"),
                "Stability": "Stable under normal conditions",
                "Reactivity Hazards": "May react with incompatible materials"
            },
            "data_sources": ["Chemical stability guidelines"],
            "notes": ["Stability assessment based on chemical structure"]
        }
        return section
    
    def _build_section_11(self, safety_data, toxicity_data):
        """Section 11: Toxicological Information"""
        tox_data = safety_data.get("toxicological", {})
        
        section = {
            "title": self.section_names[11],
            "data": {
                "Acute Toxicity": {
                    "Oral (LD50)": tox_data.get("LD50 Oral", toxicity_data.get("ld50", "Not determined")),
                    "Dermal (LD50)": tox_data.get("LD50 Dermal", "Not determined"),
                    "Inhalation (LC50)": tox_data.get("LC50 Inhalation", toxicity_data.get("lc50_inhalation_rat", "Not determined"))
                },
                "Skin Corrosion/Irritation": tox_data.get("Skin Corrosion", "May cause skin irritation"),
                "Serious Eye Damage/Irritation": tox_data.get("Serious Eye Damage", "May cause eye irritation"),
                "Respiratory Sensitization": tox_data.get("Respiratory Sensitization", "Not determined"),
                "Skin Sensitization": tox_data.get("Skin Sensitization", "Not determined"),
                "Germ Cell Mutagenicity": tox_data.get("Germ Cell Mutagenicity", "Not determined"),
                "Carcinogenicity": tox_data.get("Carcinogenicity", "Not classified"),
                "Reproductive Toxicity": tox_data.get("Reproductive Toxicity", "Not determined"),
                "STOT-Single Exposure": tox_data.get("STOT Single Exposure", "Not classified"),
                "STOT-Repeated Exposure": tox_data.get("STOT Repeated Exposure", "Not classified"),
                "Aspiration Hazard": tox_data.get("Aspiration Hazard", "Not determined"),
                "Routes of Exposure": "Inhalation, dermal contact, eye contact, ingestion",
                "Target Organs": ", ".join(toxicity_data.get("target_organs", ["Not specified"])),
                "Symptoms": "Irritation, nausea, dizziness, headache",
                "Toxicity Classification": toxicity_data.get("toxicity_class", "Class IV (Low)"),
                "Chronic Effects": "May cause organ damage with prolonged exposure",
                "Hazard Endpoints": ", ".join(toxicity_data.get("hazard_endpoints", ["None predicted"]))
            },
            "data_sources": ["Toxicity predictions", "Structure-activity relationships"],
            "notes": ["Toxicity data based on computational predictions"]
        }
        return section
    
    def _build_section_12(self, safety_data, basic_data):
        """Section 12: Ecological Information"""
        eco_data = safety_data.get("ecological", {})
        logp = basic_data.get("logp", 0)
        
        section = {
            "title": self.section_names[12],
            "data": {
                "Ecotoxicity": eco_data.get("Ecotoxicity", "May be harmful to aquatic organisms"),
                "Acute Aquatic Toxicity": {
                    "Fish LC50": eco_data.get("LC50 Fish", "Not determined"),
                    "Daphnia EC50": eco_data.get("EC50 Daphnia", "Not determined"),
                    "Algae EC50": eco_data.get("EC50 Algae", "Not determined")
                },
                "Persistence and Degradability": eco_data.get("Persistence", "Expected to be biodegradable"),
                "Biodegradability": eco_data.get("Biodegradability", "Expected to biodegrade"),
                "Bioaccumulative Potential": f"{'High' if logp > 3.5 else 'Low'} bioaccumulation potential (log P = {logp})",
                "Mobility in Soil": eco_data.get("Mobility in Soil", 
                    "Mobile" if logp < 2 else "Moderately mobile" if logp < 4 else "Low mobility"),
                "Other Adverse Effects": eco_data.get("Other Adverse Effects", 
                    "May cause long-term adverse effects in aquatic environment"),
                "Environmental Fate": "Expected to partition between water, sediment, and biota",
                "Aquatic Toxicity": "May be toxic to aquatic life",
                "Terrestrial Toxicity": "Limited data available"
            },
            "data_sources": ["Ecological modeling", "Property-based predictions"],
            "notes": ["Environmental fate based on physicochemical properties"]
        }
        return section
    
    def _build_section_13(self, safety_data):
        """Section 13: Disposal Considerations"""
        disposal_data = safety_data.get("disposal", {})
        
        section = {
            "title": self.section_names[13],
            "data": {
                "Waste Treatment Methods": disposal_data.get("Waste Treatment Methods", 
                    "Incineration at licensed hazardous waste facility"),
                "Disposal Methods": disposal_data.get("Disposal Method", 
                    "Dispose according to local, state, and federal regulations"),
                "Contaminated Packaging": disposal_data.get("Contaminated Packaging", 
                    "Containers should be completely emptied and disposed as hazardous waste"),
                "Special Precautions": "Do not dispose in regular trash or sewage system",
                "Regulatory Requirements": "Follow EPA RCRA regulations for hazardous waste disposal",
                "Recommended Method": "Contract with licensed waste disposal company",
                "Preparation for Disposal": "Collect waste in appropriate containers. Label clearly.",
                "Treatment Options": "Chemical treatment, incineration, or secure landfill",
                "Waste Code": "Consult local regulations for appropriate waste classification"
            },
            "data_sources": ["Waste disposal guidelines"],
            "notes": ["Consult local environmental regulations before disposal"]
        }
        return section
    
    def _build_section_14(self, safety_data, basic_data):
        """Section 14: Transport Information"""
        transport_data = safety_data.get("transport", {})
        is_flammable = basic_data.get("logp", 0) > 1.5
        
        section = {
            "title": self.section_names[14],
            "data": {
                "UN Number": transport_data.get("UN Number", 
                    "UN1993" if is_flammable else "Not regulated"),
                "UN Proper Shipping Name": transport_data.get("UN Proper Shipping Name", 
                    "Flammable liquid, n.o.s." if is_flammable else "Research chemical"),
                "Transport Hazard Class": transport_data.get("Transport Hazard Class", 
                    "3" if is_flammable else "Not applicable"),
                "Packing Group": transport_data.get("Packing Group", 
                    "III" if is_flammable else "Not applicable"),
                "Environmental Hazards": transport_data.get("Environmental Hazards", "Not classified"),
                "Marine Pollutant": transport_data.get("Marine Pollutant", "No"),
                "Special Precautions": transport_data.get("Special Precautions", 
                    "Follow DOT regulations for hazardous materials"),
                "Transport by Road/Rail": "Follow ADR/RID regulations where applicable",
                "Transport by Sea": "Follow IMDG Code where applicable", 
                "Transport by Air": "Follow IATA regulations where applicable",
                "Emergency Response": "Carry appropriate emergency response information"
            },
            "data_sources": ["DOT regulations", "Transport guidelines"],
            "notes": ["Verify current transport regulations before shipping"]
        }
        return section
    
    def _build_section_15(self, safety_data):
        """Section 15: Regulatory Information"""
        reg_data = safety_data.get("regulatory", {})
        
        section = {
            "title": self.section_names[15],
            "data": {
                "Safety, Health and Environmental Regulations": {
                    "TSCA Status": reg_data.get("TSCA", "Not listed"),
                    "DSL/NDSL (Canada)": reg_data.get("DSL/NDSL", "Not determined"),
                    "EINECS/ELINCS (EU)": reg_data.get("EINECS/ELINCS", "Not listed"),
                    "ENCS (Japan)": reg_data.get("ENCS", "Not determined"),
                    "IECSC (China)": reg_data.get("IECSC", "Not determined"),
                    "KECL (Korea)": reg_data.get("KECL", "Not determined"),
                    "PICCS (Philippines)": reg_data.get("PICCS", "Not determined"),
                    "AICS (Australia)": reg_data.get("AICS", "Not determined")
                },
                "WHMIS Classification": reg_data.get("WHMIS", "Not classified"),
                "GHS Classification": reg_data.get("GHS Classification", "See Section 3"),
                "SARA Title III": {
                    "Section 302 EHS": "Not listed",
                    "Section 311/312 Categories": "Not listed",
                    "Section 313 TRI": reg_data.get("SARA 313", "Not listed")
                },
                "California Proposition 65": reg_data.get("California Proposition 65", "Not listed"),
                "RCRA Hazardous Waste": "Not listed",
                "CERCLA Reportable Quantity": "Not established",
                "State Regulations": "May be regulated under state chemical laws",
                "International Regulations": "Subject to country-specific chemical regulations"
            },
            "data_sources": ["Regulatory databases"],
            "notes": ["Regulatory status may change. Verify current requirements."]
        }
        return section
    
    def _build_section_16(self, data):
        """Section 16: Other Information"""
        data_sources_used = data.get("data_sources", [])
        errors = data.get("errors", [])
        
        section = {
            "title": self.section_names[16],
            "data": {
                "Date of Preparation": datetime.now().strftime("%Y-%m-%d"),
                "Date of Last Revision": datetime.now().strftime("%Y-%m-%d"),
                "Revision Number": "1.0",
                "Prepared By": "Automated SDS Generator v3.0",
                "Data Sources Used": ", ".join(data_sources_used) if data_sources_used else "Computational predictions",
                "References": {
                    "PubChem Database": f"CID: {data.get('basic_data', {}).get('cid', 'Unknown')}",
                    "RDKit": "Open-source cheminformatics toolkit",
                    "Toxicity Predictions": "Structure-activity relationship models",
                    "Regulatory Guidelines": "GHS, OSHA, EPA standards"
                },
                "Key Literature": "Consult PubChem and peer-reviewed sources for additional data",
                "Abbreviations": {
                    "ACGIH": "American Conference of Governmental Industrial Hygienists",
                    "CAS": "Chemical Abstracts Service", 
                    "DOT": "Department of Transportation",
                    "EPA": "Environmental Protection Agency",
                    "GHS": "Globally Harmonized System",
                    "NIOSH": "National Institute for Occupational Safety and Health",
                    "OSHA": "Occupational Safety and Health Administration",
                    "PEL": "Permissible Exposure Limit",
                    "PPE": "Personal Protective Equipment",
                    "SARA": "Superfund Amendments and Reauthorization Act",
                    "SDS": "Safety Data Sheet",
                    "STEL": "Short Term Exposure Limit",
                    "TLV": "Threshold Limit Value",
                    "TSCA": "Toxic Substances Control Act",
                    "TWA": "Time Weighted Average"
                },
                "Version History": "Initial automated generation",
                "Quality Assurance": "Generated using validated computational methods",
                "Data Limitations": "Some values are computationally predicted. Laboratory verification recommended.",
                "Disclaimer": "This SDS is generated for research purposes using computational methods. Users must verify all information through laboratory testing and consult authoritative sources. No warranty is provided for accuracy or completeness.",
                "Training Information": "Ensure personnel are trained in chemical safety before use",
                "Emergency Information": "Maintain emergency contact information and procedures",
                "Update Schedule": "Review annually or when new data becomes available"
            },
            "data_sources": ["System metadata"],
            "notes": errors if errors else ["SDS generated successfully"]
        }
        return section
    
    def generate_docx_report(self, sds, compound_name="Unknown Compound"):
        """
        Generate comprehensive DOCX report from SDS data.
        Returns BytesIO buffer for Flask send_file() compatibility.
        """
        logger.info(f"[DOCX Generator] Creating Word document for {compound_name}")
        
        # Create document
        doc = Document()
        
        # Set document margins
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1) 
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
        
        # Document title and header
        title = doc.add_heading('Safety Data Sheet (SDS)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"Chemical: {compound_name}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(14)
        
        # Generation info
        generated_info = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        generated_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_info.runs[0].font.size = Pt(10)
        generated_info.runs[0].italic = True
        
        doc.add_paragraph()  # Spacing
        
        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        toc_table = doc.add_table(rows=0, cols=2)
        toc_table.style = 'Light List Accent 1'
        
        for i in range(1, 17):
            section_title = self.section_names[i]
            row = toc_table.add_row()
            row.cells[0].text = f"Section {i}"
            row.cells[1].text = section_title
        
        doc.add_page_break()
        
        # Generate all sections
        for i in range(1, 17):
            section_key = f"Section{i}"
            section_data = sds.get(section_key, {})
            section_title = section_data.get("title", f"Section {i}")
            
            # Section heading
            heading = doc.add_heading(f"{i}. {section_title}", level=1)
            
            # Section data
            data = section_data.get("data", {})
            if not data:
                doc.add_paragraph("No data available for this section.")
            else:
                # Create data table
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                
                # Set column widths
                for column in table.columns:
                    for cell in column.cells:
                        cell.width = Inches(3.0)
                
                # Add data rows
                for key, value in data.items():
                    row = table.add_row()
                    
                    # Key cell (bold)
                    key_cell = row.cells[0]
                    key_paragraph = key_cell.paragraphs[0]
                    key_run = key_paragraph.add_run(str(key))
                    key_run.bold = True
                    key_run.font.size = Pt(10)
                    
                    # Value cell
                    value_cell = row.cells[1]
                    value_paragraph = value_cell.paragraphs[0]
                    
                    # Format value based on type
                    if isinstance(value, dict):
                        # Handle nested dictionaries
                        value_text = ""
                        for sub_key, sub_value in value.items():
                            value_text += f"{sub_key}: {sub_value}\n"
                        value_text = value_text.strip()
                    elif isinstance(value, list):
                        value_text = ", ".join(str(v) for v in value if v) or "Not available"
                    elif value is None or value == "":
                        value_text = "Not available"
                    else:
                        value_text = str(value)
                    
                    # Truncate very long values
                    if len(value_text) > 1000:
                        value_text = value_text[:1000] + "... [truncated]"
                    
                    value_run = value_paragraph.add_run(value_text)
                    value_run.font.size = Pt(10)
            
            # Data sources
            sources = section_data.get("data_sources", [])
            if sources:
                sources_para = doc.add_paragraph()
                sources_run = sources_para.add_run(f"Data Sources: {', '.join(sources)}")
                sources_run.font.size = Pt(9)
                sources_run.italic = True
            
            # Notes
            notes = section_data.get("notes", [])
            if notes:
                notes_para = doc.add_paragraph()
                notes_run = notes_para.add_run(f"Notes: {'; '.join(notes)}")
                notes_run.font.size = Pt(9)
                notes_run.italic = True
            
            doc.add_paragraph()  # Section spacing
        
        # Footer with disclaimer
        doc.add_page_break()
        doc.add_heading('Important Disclaimer', level=1)
        
        disclaimer_text = """
        This Safety Data Sheet has been generated using computational methods and database information for research purposes only. 
        
        IMPORTANT WARNINGS:
        • This SDS contains predicted and estimated data that may not reflect actual chemical properties
        • All information should be verified through laboratory testing before use
        • Consult authoritative sources and conduct proper hazard assessments
        • This document does not replace professional chemical safety evaluation
        • The generators assume no responsibility for accuracy, completeness, or suitability for any purpose
        
        FOR RESEARCH USE ONLY. Not for commercial, industrial, or consumer applications.
        
        Always follow institutional safety protocols and consult with qualified safety professionals before handling any chemical substance.
        """
        
        disclaimer_para = doc.add_paragraph(disclaimer_text)
        disclaimer_run = disclaimer_para.runs[0]
        disclaimer_run.font.size = Pt(10)
        disclaimer_run.italic = True
        
        # Contact information
        doc.add_paragraph()
        contact_para = doc.add_paragraph("For questions about this SDS generation system, consult your institution's chemical safety office.")
        contact_run = contact_para.runs[0]
        contact_run.font.size = Pt(9)
        contact_run.bold = True
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"[DOCX Generator] Document created successfully for {compound_name}")
        return buffer


# ===== CONVENIENCE FUNCTIONS =====

def generate_sds_from_smiles(smiles):
    """
    Convenience function to generate complete SDS from SMILES.
    Returns SDS data structure.
    """
    generator = SDSGenerator()
    return generator.generate_comprehensive_sds(smiles)

def generate_sds_docx_from_smiles(smiles, compound_name=None):
    """
    Convenience function to generate SDS and return DOCX buffer.
    Ready for Flask send_file() usage.
    """
    generator = SDSGenerator()
    sds = generator.generate_comprehensive_sds(smiles)
    
    if not sds:
        return None
    
    # Get compound name from SDS data if not provided
    if not compound_name:
        compound_name = sds.get("Section1", {}).get("data", {}).get("Product Identifier", "Unknown Compound")
    
    return generator.generate_docx_report(sds, compound_name)

def get_sds_section_names():
    """Return dictionary mapping section numbers to names"""
    generator = SDSGenerator()
    return generator.section_names


